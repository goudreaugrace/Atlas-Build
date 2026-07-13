#!/usr/bin/env python3
"""Generate the shareable Brand Doctor System Wiki Word export.

The source of truth stays in docs/wiki plus src/data/config/wiki-nav.json.
This script converts that markdown-backed wiki into an approachable Word
reference guide for stakeholder review and team annotation.
"""

from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WIKI_DIR = ROOT / "docs" / "wiki"
NAV_PATH = ROOT / "src" / "data" / "config" / "wiki-nav.json"
EXPORT_DIR = ROOT / "public" / "exports"
DOCX_PATH = EXPORT_DIR / "bbe-brand-doctor-system-wiki.docx"

COLORS = {
    "navy": "02355A",
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "orange": "F7941D",
    "ink": "0B2545",
    "muted": "666666",
    "light_blue": "E8EEF5",
    "light_gray": "F2F4F7",
    "callout": "F4F6F9",
    "border": "D9E2EE",
    "white": "FFFFFF",
}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_spacing(style, before: int = 0, after: int = 6, line_spacing: float = 1.25) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color: str = COLORS["border"]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches: float) -> None:
    width = int(width_inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def clear_cell(cell) -> None:
    cell.text = ""


def add_runs(paragraph, text: str) -> None:
    tokens = re.split(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = rgb(COLORS["navy"])
        elif token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.color.rgb = rgb(COLORS["ink"])
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            run.font.color.rgb = rgb(COLORS["muted"])
        else:
            paragraph.add_run(token)


def markdown_blocks(content: str) -> Iterable[tuple[str, str]]:
    lines = content.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            yield f"h{len(heading.group(1))}", heading.group(2).strip()
            index += 1
            continue

        if re.match(r"^[-*]\s+", line):
            while index < len(lines) and re.match(r"^[-*]\s+", lines[index].strip()):
                yield "bullet", re.sub(r"^[-*]\s+", "", lines[index].strip())
                index += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            while index < len(lines) and re.match(r"^\d+\.\s+", lines[index].strip()):
                yield "number", re.sub(r"^\d+\.\s+", "", lines[index].strip())
                index += 1
            continue

        paragraph_lines = [line]
        index += 1
        while (
            index < len(lines)
            and lines[index].strip()
            and not re.match(r"^(#{1,3})\s+", lines[index].strip())
            and not re.match(r"^[-*]\s+", lines[index].strip())
            and not re.match(r"^\d+\.\s+", lines[index].strip())
        ):
            paragraph_lines.append(lines[index].strip())
            index += 1
        yield "p", " ".join(paragraph_lines)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(COLORS["ink"])
    set_spacing(normal, after=6, line_spacing=1.25)

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(25)
    title.font.bold = True
    title.font.color.rgb = rgb(COLORS["navy"])
    set_spacing(title, after=8, line_spacing=1.1)

    for name, size, color, before, after in (
        ("Heading 1", 16, COLORS["blue"], 18, 10),
        ("Heading 2", 13, COLORS["blue"], 14, 7),
        ("Heading 3", 12, COLORS["dark_blue"], 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        set_spacing(style, before=before, after=after, line_spacing=1.25)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = rgb(COLORS["ink"])
        fmt = style.paragraph_format
        fmt.left_indent = Inches(0.375)
        fmt.first_line_indent = Inches(-0.188)
        fmt.space_after = Pt(4)
        fmt.line_spacing = 1.25


def set_headers_and_footers(doc: Document) -> None:
    for section in doc.sections:
        header_p = section.header.paragraphs[0]
        header_p.text = ""
        header_run = header_p.add_run("BBE Brand Doctor System Wiki")
        header_run.font.name = "Calibri"
        header_run.font.size = Pt(9)
        header_run.font.bold = True
        header_run.font.color.rgb = rgb(COLORS["navy"])

        footer_p = section.footer.paragraphs[0]
        footer_p.text = ""
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_p.add_run("JSON-first prototype  |  No Magic principle  |  Generated from docs/wiki")
        footer_run.font.name = "Calibri"
        footer_run.font.size = Pt(8.5)
        footer_run.font.color.rgb = rgb(COLORS["muted"])


def add_metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    for row_index, (label, value) in enumerate(rows):
        row = table.rows[row_index]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(24)
        label_cell, value_cell = row.cells
        set_cell_width(label_cell, 1.181)
        set_cell_width(value_cell, 5.319)
        set_cell_shading(label_cell, COLORS["light_blue"])
        set_cell_shading(value_cell, COLORS["white"])
        clear_cell(label_cell)
        clear_cell(value_cell)
        label_p = label_cell.paragraphs[0]
        label_run = label_p.add_run(label)
        label_run.bold = True
        label_run.font.color.rgb = rgb(COLORS["navy"])
        add_runs(value_cell.paragraphs[0], value)


def add_cover(doc: Document, pages: list[dict]) -> None:
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = kicker.add_run("BBE BRAND DOCTOR")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = rgb(COLORS["orange"])

    title = doc.add_paragraph(style="Title")
    title.add_run("System Wiki Reference Guide")

    subtitle = doc.add_paragraph()
    subtitle.add_run(
        "A shareable, editable explanation of how the prototype works: data, rules, treatments, AI guardrails, personas, Live Consult, and operator practices."
    )
    subtitle.paragraph_format.space_after = Pt(14)

    generated = datetime.now().strftime("%B %-d, %Y at %-I:%M %p")
    add_metadata_table(
        doc,
        [
            ("Purpose", "Help stakeholders review the system without reading a full PRD or source code."),
            ("Source", "`docs/wiki` markdown plus `src/data/config/wiki-nav.json`."),
            ("Coverage", f"{len(pages)} wiki sections covering Brand Doctor foundations, No Magic AI, and operating rules."),
            ("Generated", generated),
        ],
    )

    doc.add_paragraph()
    callout = doc.add_paragraph()
    set_paragraph_shading(callout, COLORS["callout"])
    callout.paragraph_format.space_before = Pt(8)
    callout.paragraph_format.space_after = Pt(8)
    callout.add_run("No Magic principle: ").bold = True
    callout.add_run(
        "observed data, deterministic logic, and AI interpretation should stay visibly separate so questions become productive instead of mysterious."
    )

    doc.add_page_break()


def add_toc(doc: Document, pages: list[dict]) -> None:
    doc.add_heading("Document Map", level=1)
    doc.add_paragraph(
        "Use this as a discussion guide. The Word file is editable for comments, wording changes, and clarification requests; the PDF is for easy read-ahead sharing."
    )
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    widths = (0.5, 2.0, 4.0)
    headers = ("#", "Section", "What it answers")
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, COLORS["light_blue"])
        clear_cell(cell)
        run = cell.paragraphs[0].add_run(headers[idx])
        run.bold = True
        run.font.color.rgb = rgb(COLORS["navy"])
    for page_index, page in enumerate(pages, start=1):
        cells = table.add_row().cells
        values = (str(page_index), page["title"], page["summary"])
        for idx, value in enumerate(values):
            set_cell_width(cells[idx], widths[idx])
            clear_cell(cells[idx])
            add_runs(cells[idx].paragraphs[0], value)
    doc.add_page_break()


def add_wiki_section(doc: Document, page: dict, is_first: bool) -> None:
    if not is_first:
        doc.add_section(WD_SECTION.NEW_PAGE)

    content = (WIKI_DIR / page["file"]).read_text(encoding="utf-8")
    source_table_rows = [
        ("Source", f"`docs/wiki/{page['file']}`"),
        ("Audience read", page["summary"]),
    ]
    first_heading_rendered = False

    for block_type, value in markdown_blocks(content):
        if block_type == "h1":
            heading = doc.add_heading(value, level=1)
            first_heading_rendered = True
            add_metadata_table(doc, source_table_rows)
            doc.add_paragraph()
        elif block_type == "h2":
            doc.add_heading(value, level=2)
        elif block_type == "h3":
            doc.add_heading(value, level=3)
        elif block_type == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, value)
        elif block_type == "number":
            p = doc.add_paragraph(style="List Number")
            add_runs(p, value)
        else:
            p = doc.add_paragraph()
            add_runs(p, value)

    if not first_heading_rendered:
        doc.add_heading(page["title"], level=1)
        add_metadata_table(doc, source_table_rows)


def build_document() -> None:
    pages = json.loads(NAV_PATH.read_text(encoding="utf-8"))
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_document(doc)
    add_cover(doc, pages)
    add_toc(doc, pages)
    for index, page in enumerate(pages):
        add_wiki_section(doc, page, is_first=index == 0)
    set_headers_and_footers(doc)
    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    build_document()
